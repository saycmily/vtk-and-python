import requests
from bs4 import BeautifulSoup
import os
import docx
from docx import Document
from docx.shared import Inches

def split_text_by_img(html, imglist):#将html代码以图片list进行分割成块
  content_parts = [] #根据图标签将正文分割成N部分
  for imgtag in imglist: #imgtag属性是bs4.element.Tag 后面需要使用str()函数转换成string
    #print(imgtag)
    html = str(html)#强制转化为字符串方便split分割
    str_tmp = html.split(str(imgtag))[0] #取图片分割的前一个元素 加入 正文list部分
    content_parts.append(str_tmp)
    #print(len(arr))
    html = html.replace((str_tmp+str(imgtag)),'')#将正文第一部分及图片标签字符串 从html中替换抹掉作为下一个for循环的html
    #print(html)
  content_parts.append(html)#把最后一张图片后的html内容补上
  return content_parts


def pic_down(pic_url):#根据图片url保存图片，填写referer可伪装referer来源下载防盗链图片
  headers = { "Accept":"text/html,application/xhtml+xml,application/xml;",
              "Accept-Encoding":"gzip",
              "Accept-Language":"zh-CN,zh;q=0.8",
              "User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/87.0.4280.88 Safari/537.36"
              }
  #保存图片至本地，因为新浪图片url中，不带后缀，这里就加了jpg后缀名，否则生成的word会报错
  img_name = pic_url.split('/')[-1]+'.jpg'
  #img_name = pic_url.split('/')[-1]
  
  with open(img_name,'wb')as f:
   response = requests.get(pic_url,headers=headers).content
   f.write(response)
   f.close()
   return img_name

url = "https://mp.weixin.qq.com/s?__biz=MzU5Nzg1NDIxOQ==&mid=2247484677&idx=1&sn=59878ea2f0aa6b83d5cf652bb5557c33&chksm=fe4c5461c93bdd777041516363c144189380288bb381298d233f1d793625c05f55336b2dda2b#rd"
html = requests.get(url).content
soup = BeautifulSoup(html,'html.parser')
#post_detial = soup.find('div',{"class":"post_text"}).text
title = soup.h2.string;#取得title标签的文字内容
post_detial = soup.find('div',{"id":"img-content"})#限定需要保存到word的区域，
#F12查找到div标签id=sina_keyword_ad_area2 的区域是新浪博客的详情内容
#print(post_detial) #获取div区域html源文件


#img_tag_list=post_detial.select('a img')
img_tag_list=post_detial.select('a > img')#在详情内容中截取出图片img的标签list
print("发现图片标签个数："+str(len(img_tag_list)))
new_text = split_text_by_img(post_detial,img_tag_list)#按图片分割成段，返回list串
print("截取博客文章片段数："+str(len(new_text)))

#print('beg==============================================')
#print(new_text[1])
#print('end==============================================')

document = Document()
document.add_heading(title)#向文档里添加标题
i = 0
for part in new_text: #循环写入
  print('写入第 '+str(i+1)+' 个片段')
  #print(n)
  part='<html><body><div >'+str(part) #part是含html标签的字符串，下面使用BeautifuSoup时需要lxml格式化，所以需要加前缀
  # 使的每个part部分都更像一个网页，否则BeautifulSoup(part,'lxml')处理的时候会把第二部分开始的内容处理为空

  part_tag = BeautifulSoup(part,'lxml')# 如果不进行lxml处理，下面get_text()无法使用。
  document.add_paragraph(part_tag.get_text())#向文档里添加文字
  
  if (i < len(img_tag_list)):#写完片段后紧接写入位于此处分割的图片
    imgurl = img_tag_list[i].get('src')#新浪图片地址在real_src属性里，一般是src
    img_name =  pic_down(imgurl)#新浪图片有防盗链，需要加入referer='http://blog.sina.com.cn'
    print('写入第 '+str(i+1)+' 张图片：'+imgurl)
    #document.add_paragraph('此处需要插入图片'+img_name+imgurl)#向文档里添加文字
    document.add_picture(img_name)#向文档里添加图片
    #os.remove(img_name)#删除保存在本地的图片
  i=i+1


#document.save('图文.doc')#保存文档
document.save(title+'.doc')